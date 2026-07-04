"""Извлечение текста/Markdown из загруженных файлов по формату."""
from __future__ import annotations

import csv
import json
import re
from io import BytesIO, StringIO
from xml.dom import minidom

from mkg_ingestion.formats import detect_route
from mkg_ingestion.ocr import ocr_file

try:
    import chardet
except Exception:  # pragma: no cover
    chardet = None  # type: ignore[assignment]

try:
    import yaml
except Exception:  # pragma: no cover
    yaml = None  # type: ignore[assignment]

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None  # type: ignore[assignment,misc]

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover
    load_workbook = None  # type: ignore[assignment,misc]


def decode_text(content: bytes) -> str:
    """Декодирование текстовых файлов с автоопределением кодировки."""
    for bom, enc in ((b"\xff\xfe", "utf-16-le"), (b"\xfe\xff", "utf-16-be"), (b"\xef\xbb\xbf", "utf-8")):
        if content.startswith(bom):
            return content[len(bom) :].decode(enc, errors="replace")
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        pass
    if chardet is not None:
        guess = chardet.detect(content[:65536])
        enc = guess.get("encoding") or "utf-8"
        try:
            return content.decode(enc, errors="replace")
        except LookupError:
            pass
    return content.decode("utf-8", errors="replace")


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [r + [""] * (width - len(r)) for r in rows]
    header = norm[0]
    sep = ["---"] * width
    lines = [
        "| " + " | ".join(c.replace("|", "\\|") for c in header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for row in norm[1:]:
        lines.append("| " + " | ".join(c.replace("|", "\\|") for c in row) + " |")
    return "\n".join(lines)


def parse_csv(content: bytes) -> str:
    text = decode_text(content)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=";,|\t")
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(StringIO(text), dialect)
    rows = [[cell.strip() for cell in row] for row in reader if any(cell.strip() for cell in row)]
    if not rows:
        return text.strip()
    return _rows_to_markdown_table(rows)


def parse_json(content: bytes) -> str:
    text = decode_text(content)
    data = json.loads(text)
    pretty = json.dumps(data, ensure_ascii=False, indent=2)
    return f"```json\n{pretty}\n```"


def parse_yaml(content: bytes) -> str:
    if yaml is None:
        return decode_text(content)
    text = decode_text(content)
    data = yaml.safe_load(text)
    dumped = yaml.dump(data, allow_unicode=True, sort_keys=False)
    return f"```yaml\n{dumped}```"


def parse_xml(content: bytes) -> str:
    text = decode_text(content)
    try:
        dom = minidom.parseString(text.encode("utf-8", errors="replace"))
        pretty = dom.toprettyxml(indent="  ")
        lines = [ln for ln in pretty.splitlines() if ln.strip()]
        body = "\n".join(lines[:500])
        return f"```xml\n{body}\n```"
    except Exception:
        return f"```xml\n{text.strip()}\n```"


def _heading_level(style_name: str) -> int:
    m = re.search(r"(\d+)", style_name or "")
    if m:
        return min(int(m.group(1)), 6)
    if "Title" in (style_name or ""):
        return 1
    return 2


def parse_docx(content: bytes) -> str:
    if DocxDocument is None:
        raise RuntimeError("python-docx не установлен")
    doc = DocxDocument(BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading" in style or style == "Title":
            level = _heading_level(style)
            parts.append("#" * level + " " + text)
        else:
            parts.append(text)
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            parts.append(_rows_to_markdown_table(rows))
    if not parts:
        raise ValueError("DOCX не содержит извлекаемого текста")
    return "\n\n".join(parts)


def parse_xlsx(content: bytes) -> str:
    if load_workbook is None:
        raise RuntimeError("openpyxl не установлен")
    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c).strip() for c in row]
            if any(cells):
                rows.append(cells)
        if rows:
            parts.append(f"## {sheet_name}\n\n{_rows_to_markdown_table(rows)}")
    wb.close()
    if not parts:
        raise ValueError("XLSX не содержит данных")
    return "\n\n".join(parts)


def parse_markdown(content: bytes) -> str:
    return decode_text(content)


def parse_plain_text(content: bytes) -> str:
    return decode_text(content)


async def extract_from_bytes(file_name: str, content: bytes) -> str:
    """Маршрутизация по формату → текст/Markdown для pipeline."""
    route = detect_route(file_name)
    if route == "ocr":
        return await ocr_file(file_name, content)
    if route == "markdown":
        return parse_markdown(content)
    if route == "text":
        return parse_plain_text(content)
    if route == "csv":
        return parse_csv(content)
    if route == "json":
        return parse_json(content)
    if route == "yaml":
        return parse_yaml(content)
    if route == "xml":
        return parse_xml(content)
    if route == "docx":
        return parse_docx(content)
    if route == "xlsx":
        return parse_xlsx(content)
    raise ValueError(f"Неподдерживаемый формат: {file_name}")
